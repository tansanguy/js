from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("reports/b4_optimization_final_validation/b4_optimization_final_validation_one_page_20260608.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "000000") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(8.5)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_run(run, size=9.5, bold=False, color="000000") -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_kv_paragraph(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(label)
    style_run(r, size=9.2, bold=True, color="1F4D78")
    r = p.add_run(body)
    style_run(r, size=9.2)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(9.3)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("B4 최적화 비교 실험 및 최종 검증 결과 보고")
    style_run(r, size=17, bold=True, color="0B2545")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(6)
    r = meta.add_run("기준 BO run: s1forced_bo_n1_m50_t6_w4_proc_20260608_001950 · 최종 검증 run: final_destination_validation_bo_best_20260608 · 작성일: 2026-06-08")
    style_run(r, size=8.3, color="555555")

    callout = doc.add_table(rows=1, cols=1)
    callout.autofit = False
    callout.columns[0].width = Inches(7.0)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "E8EEF5")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("결론: ")
    style_run(r, size=10, bold=True, color="0B2545")
    r = p.add_run(
        "BO로 선택한 B4 theta는 3개 최종 경로에서 B04 대비 응급차 추가 지연을 평균 223.4s -> 99.1s로 줄였다. "
        "개선폭은 평균 124.4초, 약 55.7%이며, 실패 run·teleport·도착 실패는 없었다."
    )
    style_run(r, size=10, color="0B2545")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("핵심 결과")
    style_run(r, size=12, bold=True, color="2E74B5")

    table = doc.add_table(rows=1, cols=6)
    table.autofit = False
    widths = [1.35, 1.12, 1.12, 1.12, 1.18, 1.18]
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    headers = ["항목", "B04", "B4", "개선", "일반차 영향", "해석"]
    for idx, h in enumerate(headers):
        c = table.rows[0].cells[idx]
        set_cell_text(c, h, bold=True, color="0B2545")
        set_cell_shading(c, "F2F4F7")

    rows = [
        ("Route 006", "254.4s", "88.9s", "-165.5s", "-6.2s", "큰 개선"),
        ("Route 016", "280.2s", "125.8s", "-154.4s", "-7.1s", "큰 개선"),
        ("Route 011", "135.7s", "82.5s", "-53.2s", "+9.2s", "trade-off"),
        ("평균", "223.4s", "99.1s", "-124.4s", "-1.4s", "55.7% 감소"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text, bold=(row[0] == "평균"), color="000000")
            if row[0] == "평균":
                set_cell_shading(cells[idx], "F4F6F9")

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for m in ("top", "bottom", "start", "end"):
                node = tc_mar.find(qn(f"w:{m}"))
                if node is None:
                    node = OxmlElement(f"w:{m}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "80")
                node.set(qn("w:type"), "dxa")

    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    add_kv_paragraph(
        doc,
        "왜 개선됐나: ",
        "B4는 응급차 접근 시 Stage3 preemption을 실제로 사용했다. 평균 preemption은 route 006/016/011에서 각각 11.1/12.7/6.8회였고, Stage2 hold는 0.2회 내외라 개선의 주된 원인은 적극적인 녹색 제공/우선 신호 개입이다.",
    )
    add_kv_paragraph(
        doc,
        "최적화 비교: ",
        "단일 seed 최종 score는 BO 91.20, CMA-ES 90.04, Random Search 90.34이다. 차이가 작고 seed 1판이므로 우열 결론보다 'BO는 초반에 빠르게 좋은 영역에 도달했고, 세 방법 모두 유사한 영역에 수렴했다'로 해석해야 한다.",
    )
    add_kv_paragraph(
        doc,
        "가중치 민감도: ",
        "1:1, 5:1, 10:1, 15:1, 20:1 모두 동일 후보(t_lead=31, delta_T_thr=28, G_ext=30, Q_ratio=0.19, tau=0.81)를 선택했다. 이는 10:1이 유일한 정답이라는 뜻이 아니라, tested weight range에서 후보가 robust하다는 뜻이다.",
    )
    add_kv_paragraph(
        doc,
        "의사결정 의미: ",
        "B4는 응급차 지연 감소 효과가 크고 검증 안정성도 좋다. 다만 route 011에서는 일반차 평균 통행시간이 9.2초 증가했으므로, 최종 발표에서는 응급차 개선과 일반차 비용을 함께 제시하는 Pareto/trade-off 메시지가 필요하다.",
    )

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(2)
    foot.paragraph_format.space_after = Pt(0)
    r = foot.add_run(
        "근거 데이터: table1_best_so_far.csv, table3_pareto.csv, final/selected_mode_averages.csv. "
        "GP 그림은 실제 재평가가 아닌 BO 관측값 기반 surrogate 1D slice이다."
    )
    style_run(r, size=7.8, color="666666")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
